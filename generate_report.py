try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError:
    print("python-docx is not installed. Please run pip install python-docx")
    exit(1)

doc = Document()
doc.add_heading('KAVACH-X Benchmark Testing Report', 0)

doc.add_heading('1. Overview', level=1)
doc.add_paragraph('This document summarizes the testing results of the KAVACH-X OpenEnv benchmark, including local sanity checks, AI agent baseline testing, and the specific bug fixes implemented in inference.py.')

doc.add_heading('2. Phase 1 & 2: Environment Logic & Random Agent', level=1)
doc.add_paragraph('Test Conducted: Manually ran actions through KavachXEnv ("test_env.py") and simulated a purely Random Agent across 100 days/iterations.')
p = doc.add_paragraph()
p.add_run('Result: ').bold = True
p.add_run('Passed. The environment correctly tracks days iteratively, handles strict action constraints, properly subtracts unit budget, and truncates episodes. The Random Agent baseline failed as expected, receiving an appropriately low score (0.00), proving grader integrity.')

doc.add_heading('3. Phase 4: Baseline Inference & Grading Logic', level=1)
doc.add_paragraph('Test Conducted: Evaluated all 3 scenarios (Easy, Medium, Hard) using the baseline inference fallback heuristic in "inference.py".')
p = doc.add_paragraph()
p.add_run('Scores Obtained:\n').bold = True
p.add_run('- Easy Scenario: 0.710\n- Medium Scenario: 0.690\n- Hard Scenario: 0.233\n\n')
p.add_run('Analysis: ').bold = True
p.add_run('The OpenEnv grading mechanics successfully differentiate scenarios by difficulty. The Easy scenario allows high performance natively, while the Hard scenario significantly suppresses the agent score by demanding cross-domain intelligence and hiding traps/decoys.')

doc.add_heading('4. Changes Made to inference.py', level=1)
doc.add_paragraph('During initial testing, "inference.py" crashed with a "TypeError: \'<\' not supported between instances of dict and dict". The issue occurred because the LLM heuristic attempted to sort a nested dictionary representing the belief states instead of explicitly sorting by the numeric metric "fraud_prob". I patched the code securely.')

doc.add_heading('Fix 1: Top Belief Output String', level=2)
doc.add_paragraph('Old Code:\n> top_beliefs = sorted(beliefs.items(), key=lambda x: x[1], reverse=True)[:5]')
doc.add_paragraph('New Code:\n> top_beliefs = sorted(beliefs.items(), key=lambda x: x[1].get("fraud_prob", 0.0), reverse=True)[:5]')

doc.add_heading('Fix 2: Suspect Filter Array Generation', level=2)
doc.add_paragraph('Old Code:\n> sorted(beliefs.items(), key=lambda x: x[1], reverse=True)')
doc.add_paragraph('New Code:\n> sorted(beliefs.items(), key=lambda x: x[1].get("fraud_prob", 0.0), reverse=True)')

doc.add_heading('Fix 3: Top Probability Extraction', level=2)
doc.add_paragraph('Old Code:\n> top_prob = beliefs.get(top_suspects[0], 0.0)')
doc.add_paragraph('New Code:\n> top_prob = beliefs.get(top_suspects[0], {}).get("fraud_prob", 0.0)')

doc.add_heading('5. System Output Formatting Check', level=1)
doc.add_paragraph('OpenEnv Logging Validation: ')
doc.add_paragraph('The inference script accurately yields the rigorous OpenEnv terminal output format [START, STEP, END] flawlessly:')
doc.add_paragraph('[START] task=kavach-easy env=kavach-x model=Qwen/Qwen2.5-72B-Instruct\n[STEP] step=1 action=FLAG_SUSPICIOUS(A-001) reward=0.00 done=false error=null\n[END] success=true steps=5 score=0.710 rewards=0.00,0.00,0.00,0.00,0.71')

file_path = 'KAVACH-X_Testing_Report.docx'
doc.save(file_path)
print(f"Successfully generated Word Document at: {file_path}")
